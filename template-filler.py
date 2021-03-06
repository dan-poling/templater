
def template_filler(template_filename, sheet_filename, result_filename):
    import re
    from docx import Document
    import pandas as pd

    # read template .docx
    doc = Document(template_filename)

    # pandas to read an .xlsx as a DataFrame, then turn it into a dictionary 
    d = pd.read_excel(sheet_filename, index_col=0).to_dict()

    # This is the main function. It takes the doc, regex object, replacement string.
    def docx_replace_regex(doc_obj, regex, replace):

        # This looks without disturbing formats
        for p in doc_obj.paragraphs:
            if regex.search(p.text):
                inline = p.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if regex.search(inline[i].text):
                        text = regex.sub(replace, inline[i].text)
                        inline[i].text = text

        # This looks inside of any tables
        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    docx_replace_regex(cell, regex, replace)

    # can use this dictionary to replace a key with its value
    for key, value in d[2].items():
        key_re = re.compile(key)
        docx_replace_regex(doc, key_re, value)

    # Save our results as a new Word Document
    doc.save(result_filename)
